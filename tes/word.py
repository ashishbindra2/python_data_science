from docx import Document
from googletrans import Translator

# Function to read the Word document
def read_word_file(file_path):
    doc = Document(file_path)
    full_text = []
    
    # Read paragraphs
    for para in doc.paragraphs:
        full_text.append(para.text)
    
    # Read tables
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text for cell in row.cells]
            full_text.append('\t'.join(row_text))
    
    return full_text

# Function to translate the paragraphs
def translate_paragraphs(paragraphs, dest_language='es'):
    translator = Translator()
    translated_paragraphs = []
    for para in paragraphs:
        if para.strip():  # Translate non-empty paragraphs
            translation = translator.translate(para, dest=dest_language)
            translated_paragraphs.append(translation.text)
        else:
            translated_paragraphs.append('')  # Preserve empty paragraphs
    return translated_paragraphs

# Function to write the translated paragraphs to a new Word document
def write_word_file(translated_paragraphs, output_path):
    doc = Document()
    for para in translated_paragraphs:
        doc.add_paragraph(para)
    doc.save(output_path)

# File paths
file_path = 'Holidays.docx'
output_path = 'Translated_Holidays.docx'

# Read the Word document
paragraphs = read_word_file(file_path)
print(paragraphs)
# Translate the paragraphs
translated_paragraphs = translate_paragraphs(paragraphs, dest_language='es')

# Write the translated paragraphs to a new Word document
write_word_file(translated_paragraphs, output_path)

print(f'Translated document saved to: {output_path}')
