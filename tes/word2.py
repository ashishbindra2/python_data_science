from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt
from googletrans import Translator

# Function to read the Word document and translate its content
def translate_document(input_path, output_path, dest_language='hi'):
    translator = Translator()
    doc = Document(input_path)
    new_doc = Document()

    # Translate paragraphs
    for para in doc.paragraphs:
        if para.text.strip():  # Translate non-empty paragraphs
            translation = translator.translate(para.text, dest=dest_language)
            new_para = new_doc.add_paragraph()
            for run in para.runs:
                new_run = new_para.add_run(translation.text if run.text.strip() else run.text)
                new_run.bold = run.bold
                new_run.italic = run.italic
                new_run.underline = run.underline
                new_run.font.name = run.font.name
                new_run.font.size = run.font.size
        else:
            new_doc.add_paragraph(para.text)  # Preserve empty paragraphs

    # Translate tables
    for table in doc.tables:
        new_table = new_doc.add_table(rows=len(table.rows), cols=len(table.columns))
        for i, row in enumerate(table.rows):
            for j, cell in enumerate(row.cells):
                if cell.text.strip():  # Translate non-empty cells
                    translation = translator.translate(cell.text, dest=dest_language)
                    new_table.cell(i, j).text = translation.text
                else:
                    new_table.cell(i, j).text = cell.text  # Preserve empty cells

    # Preserve images
    for rel in doc.part.rels.values():
        if "image" in rel.target_ref:
            new_doc.part.rels[rel.rId] = rel

    new_doc.save(output_path)

# File paths
input_path = 'Holidays.docx'
output_path = 'Translated_Holidays.docx'

# Translate the document
translate_document(input_path, output_path, dest_language='hi')

print(f'Translated document saved to: {output_path}')
