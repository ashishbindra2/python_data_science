from docx import Document
from googletrans import Translator

# Load the DOCX file /mnt/data/
doc = Document('Holidays.docx')

# Initialize the translator
translator = Translator()

# Function to translate text
def translate_text(doc, target_language):
    # Translate paragraphs
    for para in doc.paragraphs:
        if para.text.strip():  # Check if paragraph is not empty
            translated_text = translator.translate(para.text, dest=target_language).text
            para.text = translated_text

    # Translate text in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():  # Check if cell is not empty
                    translated_text = translator.translate(cell.text, dest=target_language).text
                    cell.text = translated_text

# Translate all text to Hindi
translate_text(doc, 'hi')

# Save the modified document
doc.save('Modified_Holidays_Hindi.docx')
